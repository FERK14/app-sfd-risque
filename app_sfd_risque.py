import streamlit as st
import pandas as pd
import openai
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document

# 🔐 Configuration API OpenAI
openai.api_key = st.secrets["openai_api_key"]

st.set_page_config(page_title="Supervision IA - Rapport GPT", layout="wide")

# 📊 Charger les données
df = pd.read_excel("donnees_SFD_simulees.xlsx")

# 🔄 Préparer les données
X = df.drop(columns=['En_risque'])
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# 🧠 Modèle IA - Random Forest
model = RandomForestClassifier(n_estimators=100, random_state=42)
model.fit(X_scaled, df['En_risque'])
df['Prediction_IA'] = model.predict(X_scaled)
df['Probabilité_Risque'] = model.predict_proba(X_scaled)[:, 1]

# ✍️ Fonction de génération GPT-3.5
def rediger_commentaire_gpt(row):
    prompt = f"""
Tu es un superviseur financier. Rédige un rapport professionnel de 150 mots sur cette institution :

- Encours de crédits : {row['Encours_credits']:,} FCFA
- Taux de créances douteuses : {row['Creances_douteuses']:.2f} %
- Ratio de liquidité : {row['Ratio_liquidite']:.2f} %
- Ratio de solvabilité : {row['Ratio_solvabilite']:.2f} %
- Nombre d'agences : {int(row['Nb_agences'])}
- Rendement des actifs : {row['Rendement_actifs']:.2f} %
- Risque IA estimé : {row['Probabilité_Risque']:.2%}

Analyse les forces, les faiblesses et propose une recommandation claire à l’attention des superviseurs.
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Erreur OpenAI : {e}"

# 🖥️ Interface Streamlit
st.title("📊 Supervision des SFD - IA & Rapport GPT")

# SECTION 1 : Analyse individuelle
st.header("🔍 Analyse d’un SFD")
sfd_index = st.selectbox("Sélectionnez un SFD", df.index)
sfd_data = df.loc[sfd_index]

st.subheader("📌 Indicateurs financiers")
col1, col2, col3 = st.columns(3)
col1.metric("Encours crédits (FCFA)", f"{sfd_data['Encours_credits']:,}")
col2.metric("Créances douteuses (%)", f"{sfd_data['Creances_douteuses']:.2f}")
col3.metric("Ratio liquidité (%)", f"{sfd_data['Ratio_liquidite']:.2f}")

col4, col5, col6 = st.columns(3)
col4.metric("Solvabilité (%)", f"{sfd_data['Ratio_solvabilite']:.2f}")
col5.metric("Agences", int(sfd_data['Nb_agences']))
col6.metric("Rendement actifs (%)", f"{sfd_data['Rendement_actifs']:.2f}")

st.subheader("🧠 Prédiction IA")
proba = sfd_data['Probabilité_Risque']
if sfd_data['Prediction_IA'] == 1:
    st.error(f"⚠️ Risque détecté (probabilité : {proba:.2%})")
else:
    st.success(f"✅ Pas de risque détecté (probabilité : {proba:.2%})")

st.subheader("📈 Ratio de liquidité - Vue comparative")
fig, ax = plt.subplots(figsize=(10, 6))
sns.barplot(x=df['Ratio_liquidite'], y=df.index, orient='h', ax=ax, palette="coolwarm")
ax.axvline(x=100, color='red', linestyle='--')
ax.set_xlabel("Ratio de liquidité (%)")
ax.set_ylabel("SFD")
st.pyplot(fig)

# SECTION 2 : Rapport Word avec GPT
st.header("📝 Rapport Word - Top 10 SFD à risque")

if st.button("Générer le rapport avec GPT-3.5"):
    top_10 = df.sort_values(by='Probabilité_Risque', ascending=False).head(10)
    doc = Document()
    doc.add_heading('Rapport de supervision - Top 10 SFD à risque', 0)

    with st.spinner("⏳ Rédaction des commentaires par GPT..."):
        for idx, row in top_10.iterrows():
            commentaire = rediger_commentaire_gpt(row)
            doc.add_heading(f"SFD {idx+1}", level=1)
            doc.add_paragraph(f"Encours : {row['Encours_credits']:,} FCFA")
            doc.add_paragraph(f"Score IA (risque) : {row['Probabilité_Risque']:.2%}")
            doc.add_paragraph("📝 Commentaire généré par GPT :")
            doc.add_paragraph(commentaire)

    rapport_path = "rapport_top_10_gpt.docx"
    doc.save(rapport_path)

    with open(rapport_path, "rb") as f:
        st.download_button("📥 Télécharger le rapport Word", f, file_name=rapport_path)
